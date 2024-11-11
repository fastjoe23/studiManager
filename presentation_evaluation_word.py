from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EvaluationCreationError(Exception):
    pass


class PresentationEvaluationDocx:
    def __init__(
        self,
        save_path,
        evaluation_type,
        course_name,
        topic,
        lecturers,
        student,
        date,
        time,
    ):
        self.save_path = save_path
        self.topic = topic
        self.course_name = course_name
        self.lecturers = lecturers
        self.student_name = student
        self.evaluation_type = evaluation_type
        self.date = date
        self.time = time

    def create_evaluation(self):
        try:
            doc = Document()
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(1.2)  # Adjust top margin as needed

            # Set title
            if self.evaluation_type == "Bachelorarbeit":
                abbrevation_type = "BA"
            elif self.evaluation_type == "Projektarbeit 1":
                abbrevation_type = "PA 1"
            elif self.evaluation_type == "Projektarbeit 2":
                abbrevation_type = "PA 2"
            else:
                abbrevation_type = self.evaluation_type

            title = f"Bewertung mündliche {abbrevation_type}-Prüfung Kurs {self.course_name}"
            heading = doc.add_heading(level=2)
            run = heading.add_run(title)
            run.bold = True
            run.font.size = Pt(13)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            #  Function for adding chapter titles
            def add_chapter_title(doc, title):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(title)
                run.bold = True
                run.underline = True
                run.font.size = Pt(12)
                paragraph.paragraph_format.space_after = Pt(1)

            # Add details in a single paragraph
            details = (
                f"Titel: {self.topic}\n"
                f"Vortragende/r: {self.student_name}\n"
                f"Datum: {self.date}, {self.time} Uhr\n"
                f"Gutachter*in: {', '.join(str(x) for x in self.lecturers)}\n"
                "Gesund/Prüfungsfähig:  O  Ja     O Nein"
            )
            paragraph = doc.add_paragraph(details)
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.bold = True

            add_chapter_title(doc, "(A) Präsentation")

            # Add tables
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Vortragsstil\n(Umfang und Rhetorik)"
            # hdr_cells[1].text = ""
            hdr_cells[1].text = "spannend / überzeugend"
            hdr_cells[2].text = "O---O---O---O---O"
            hdr_cells[3].text = "langatmig/zäh und wenig überzeugend"
            # Adjust paragraph spacing
            paragraph = table.rows[0].cells[0].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(2)
            paragraph_format.space_after = Pt(2)

            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Fachkenntnisse\n(Qualität + Logik)"
            hdr_cells[1].text = "kompetenter Eindruck, durchgehend logisch"
            hdr_cells[2].text = "O---O---O---O---O"
            hdr_cells[3].text = "kein kompetenter Eindruck, roter Faden fehlt"
            # Adjust paragraph spacing
            paragraph = table.rows[0].cells[0].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(2)
            paragraph_format.space_after = Pt(2)

            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Unterlagen\n(Folien + Handout)"
            hdr_cells[1].text = "gut aufbereitet, gut verständlich"
            hdr_cells[2].text = "O---O---O---O---O"
            hdr_cells[3].text = "schlecht aufbereitet, wenig verständlich"
            # Adjust paragraph spacing
            paragraph = table.rows[0].cells[0].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(2)
            paragraph_format.space_after = Pt(2)

            add_chapter_title(doc, "(B) Diskussion")

            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Antwortverhalten"
            hdr_cells[1].text = "kompetente Antworten"
            hdr_cells[2].text = "O---O---O---O---O"
            hdr_cells[3].text = "wenig überzeugend"

            add_chapter_title(doc, "(C) Allgemeine Anmerkungen:")
            doc.add_paragraph("-")
            doc.add_paragraph("-")
            doc.add_paragraph("-")

            doc.add_paragraph("\n" * 8)
            add_chapter_title(doc, "(D) Note:")
            # Add signatures
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Unterschriften:  ")
            run.font.size = Pt(12)
            run.bold = True
            paragraph.add_run().add_picture(
                "C:/Users/Offtermatt/OneDrive - Duale Hochschule Baden-Württemberg Stuttgart/Dokumente/Verwaltung/Unterschrift.jpg",
                width=Cm(3),
            )
            paragraph.add_run().add_picture(
                "C:/Users/Offtermatt/OneDrive - Duale Hochschule Baden-Württemberg Stuttgart/Dokumente/Verwaltung/TK.jpg",
                width=Cm(3),
            )

            doc_file_path = (
                self.save_path
                + f"/{str(self.course_name)}_{str(self.evaluation_type).replace(' ', '')}_{str(self.student_name).replace(' ', '')}_Bewertungsbogen.docx"
            )
            doc.save(doc_file_path)
            return doc_file_path

        except Exception as e:
            error_message = f"Fehler bei der Erstellung der Bewertung: {self.topic} {e}"
            raise EvaluationCreationError(error_message) from e
